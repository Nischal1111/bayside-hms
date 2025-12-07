#!/usr/bin/env python3
"""
Script to add comprehensive user manual sections to Final.docx
Organized by student contributions
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to the document"""
    p = doc.add_paragraph(text)
    if bold and len(p.runs) > 0:
        p.runs[0].bold = True
    if italic and len(p.runs) > 0:
        p.runs[0].italic = True
    return p

def add_bullet(doc, text):
    """Add a bullet point"""
    p = doc.add_paragraph()
    p.add_run("• " + text)
    return p

def add_number(doc, num, text):
    """Add a numbered point"""
    p = doc.add_paragraph()
    p.add_run(f"{num}. {text}")
    return p

def add_code_block(doc, code):
    """Add a code block to the document"""
    p = doc.add_paragraph(code)
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def main():
    # Open existing document
    print("Opening Final.docx...")
    doc = Document('Final.docx')

    print("Adding new sections...")

    # Add page break before new sections
    doc.add_page_break()

    # ==========================================
    # HARDWARE AND SOFTWARE SETUP
    # By: Krischal Jung Thakuri
    # ==========================================
    add_heading(doc, "3. HARDWARE AND SOFTWARE SETUP", level=1)
    add_paragraph(doc, "Prepared by: Krischal Jung Thakuri (Backend/Database Developer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "3.1 Hardware Requirements", level=2)

    add_heading(doc, "3.1.1 Development Environment", level=3)
    add_paragraph(doc, "Minimum Requirements:")
    add_bullet(doc, "Processor: Intel Core i3 / AMD Ryzen 3 or equivalent")
    add_bullet(doc, "RAM: 4GB minimum (8GB recommended)")
    add_bullet(doc, "Storage: 10GB free disk space")
    add_bullet(doc, "Network: Stable internet connection")
    doc.add_paragraph()

    add_heading(doc, "3.1.2 Production Server Requirements", level=3)
    add_bullet(doc, "Processor: Intel Core i5 / AMD Ryzen 5 or better (quad-core)")
    add_bullet(doc, "RAM: 8GB minimum (16GB recommended for high traffic)")
    add_bullet(doc, "Storage: 50GB SSD (for database and application files)")
    add_bullet(doc, "Network: 100 Mbps internet connection with static IP")
    add_bullet(doc, "Backup Storage: External or cloud backup solution")
    doc.add_paragraph()

    add_heading(doc, "3.2 Software Prerequisites", level=2)

    add_heading(doc, "3.2.1 Operating System", level=3)
    add_paragraph(doc, "The system supports multiple operating systems:")
    add_bullet(doc, "Windows 10/11")
    add_bullet(doc, "macOS 10.15 or higher")
    add_bullet(doc, "Linux (Ubuntu 20.04+ recommended)")
    doc.add_paragraph()

    add_heading(doc, "3.2.2 Node.js and npm", level=3)
    add_paragraph(doc, "Node.js version 18.0.0 or higher is required.")
    add_code_block(doc, "node --version  # Should show v18.0.0 or higher\nnpm --version   # Should show 9.0.0 or higher")
    doc.add_paragraph()

    add_heading(doc, "3.2.3 MySQL Database Server", level=3)
    add_paragraph(doc, "MySQL 8.0 or higher is required for the database backend.")
    doc.add_paragraph()
    add_paragraph(doc, "Windows Installation:", bold=True)
    add_bullet(doc, "Download MySQL Installer from mysql.com")
    add_bullet(doc, "Run the installer and select 'MySQL Server' and 'MySQL Workbench'")
    add_bullet(doc, "Set root password during installation")
    doc.add_paragraph()

    add_paragraph(doc, "macOS Installation:", bold=True)
    add_code_block(doc, "brew install mysql\nbrew services start mysql\nmysql_secure_installation")
    doc.add_paragraph()

    add_paragraph(doc, "Linux (Ubuntu) Installation:", bold=True)
    add_code_block(doc, "sudo apt update\nsudo apt install mysql-server\nsudo systemctl start mysql\nsudo mysql_secure_installation")
    doc.add_paragraph()

    add_heading(doc, "3.3 Installation Steps", level=2)

    add_heading(doc, "Step 1: Clone the Repository", level=3)
    add_code_block(doc, "git clone https://github.com/Nischal1111/bayside-hms.git\ncd bayside-hms")
    doc.add_paragraph()

    add_heading(doc, "Step 2: Install Dependencies", level=3)
    add_code_block(doc, "npm install")
    add_paragraph(doc, "This installs all required packages including Next.js, React, database drivers, authentication libraries, and UI components.")
    doc.add_paragraph()

    add_heading(doc, "Step 3: Database Setup", level=3)
    add_paragraph(doc, "Create MySQL Database:", bold=True)
    add_code_block(doc, """mysql -u root -p

CREATE DATABASE bayside_hms CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;
CREATE USER 'bayside_admin'@'localhost' IDENTIFIED BY 'YourSecurePassword123!';
GRANT ALL PRIVILEGES ON bayside_hms.* TO 'bayside_admin'@'localhost';
FLUSH PRIVILEGES;
EXIT;""")
    doc.add_paragraph()

    add_paragraph(doc, "Import Database Schema:", bold=True)
    add_code_block(doc, "mysql -u bayside_admin -p bayside_hms < database/schema-mysql.sql")
    doc.add_paragraph()

    add_heading(doc, "Step 4: Environment Configuration", level=3)
    add_paragraph(doc, "Create .env.local file in the project root with the following content:")
    add_code_block(doc, """DATABASE_URL=mysql://bayside_admin:YourSecurePassword123!@localhost:3306/bayside_hms
JWT_SECRET=your_very_secure_random_32_character_secret_key
NODE_ENV=development""")
    doc.add_paragraph()

    add_heading(doc, "Step 5: Create Admin Account", level=3)
    add_code_block(doc, "npm run create-admin")
    add_paragraph(doc, "Default Admin Credentials:")
    add_bullet(doc, "Email: admin@bayside-hms.com")
    add_bullet(doc, "Password: Admin@123456")
    add_paragraph(doc, "⚠️ IMPORTANT: Change this password immediately after first login!", bold=True)
    doc.add_paragraph()

    add_heading(doc, "Step 6: Start the Development Server", level=3)
    add_code_block(doc, "npm run dev")
    add_paragraph(doc, "The application will be available at: http://localhost:3000")
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # ==========================================
    # USER MANUALS
    # By: Hillson
    # ==========================================
    add_heading(doc, "4. USER MANUAL - ADMINISTRATOR", level=1)
    add_paragraph(doc, "Prepared by: Hillson (Frontend Developer & UI/UX Designer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "4.1 Administrator Role Overview", level=2)
    add_paragraph(doc, "Administrators are the system managers with the highest level of access. Their primary responsibilities include:")
    add_bullet(doc, "Approving new doctor registrations")
    add_bullet(doc, "Managing all users (patients and doctors)")
    add_bullet(doc, "Monitoring system statistics and operations")
    add_bullet(doc, "Overseeing appointment workflows")
    add_bullet(doc, "Ensuring system integrity and security")
    doc.add_paragraph()

    add_heading(doc, "4.2 Accessing the Admin Portal", level=2)
    add_paragraph(doc, "Login Process:", bold=True)
    add_paragraph(doc, "Step 1: Navigate to Admin Portal")
    add_bullet(doc, "Open browser and go to: http://localhost:3000/admin")
    add_bullet(doc, "Or click 'Admin Portal' button on the homepage")
    doc.add_paragraph()

    add_paragraph(doc, "Step 2: Enter Credentials")
    add_bullet(doc, "Email: admin@bayside-hms.com")
    add_bullet(doc, "Password: Admin@123456 (change after first login)")
    doc.add_paragraph()

    add_paragraph(doc, "Step 3: Click 'Sign In' - You'll be redirected to the Admin Dashboard")
    doc.add_paragraph()

    add_heading(doc, "4.3 Admin Dashboard Overview", level=2)
    add_paragraph(doc, "The admin dashboard features a sidebar with the following sections:")
    add_bullet(doc, "🏠 Dashboard - Main overview with statistics")
    add_bullet(doc, "👤 Pending Approvals - New doctor registrations awaiting approval")
    add_bullet(doc, "👨‍⚕️ Manage Doctors - All approved doctors")
    add_bullet(doc, "👥 Manage Patients - All registered patients")
    doc.add_paragraph()

    add_heading(doc, "4.4 Core Administrator Functions", level=2)

    add_heading(doc, "Function 1: Approve Doctor Registrations", level=3)
    add_paragraph(doc, "Purpose: Verify and activate new doctor accounts")
    doc.add_paragraph()
    add_paragraph(doc, "Process:")
    add_number(doc, 1, "Click 'Pending Approvals' in sidebar")
    add_number(doc, 2, "Review doctor information (name, specialization, license number)")
    add_number(doc, 3, "Verify credentials")
    add_number(doc, 4, "Click green 'Approve' button")
    add_number(doc, 5, "Confirm approval in dialog")
    add_number(doc, 6, "Success message appears and doctor can now login")
    doc.add_paragraph()

    add_heading(doc, "Function 2: Manage Doctors", level=3)
    add_bullet(doc, "Click 'Manage Doctors' in sidebar")
    add_bullet(doc, "View list of all approved doctors")
    add_bullet(doc, "See doctor details: specialization, phone, email, license")
    add_bullet(doc, "Search by name, specialization, or license number")
    doc.add_paragraph()

    add_heading(doc, "Function 3: Manage Patients", level=3)
    add_bullet(doc, "Click 'Manage Patients' in sidebar")
    add_bullet(doc, "View list of all registered patients")
    add_bullet(doc, "See patient details: contact info, blood group, emergency contacts")
    add_bullet(doc, "View patient medical history and appointments")
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # USER MANUAL - DOCTOR
    add_heading(doc, "5. USER MANUAL - DOCTOR", level=1)
    add_paragraph(doc, "Prepared by: Hillson (Frontend Developer & UI/UX Designer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "5.1 Doctor Role Overview", level=2)
    add_paragraph(doc, "Doctors use the system to:")
    add_bullet(doc, "Manage patient appointments")
    add_bullet(doc, "Create and maintain electronic medical records (EMR)")
    add_bullet(doc, "View patient medical history")
    add_bullet(doc, "Prescribe medications and treatments")
    add_bullet(doc, "Receive patient feedback and ratings")
    doc.add_paragraph()

    add_heading(doc, "5.2 Doctor Registration and Activation", level=2)
    add_paragraph(doc, "Step 1: Register for Account", bold=True)
    add_bullet(doc, "Navigate to http://localhost:3000/register")
    add_bullet(doc, "Fill in personal information (name, email, password)")
    add_bullet(doc, "Fill in professional information (phone, license number, specialization)")
    add_bullet(doc, "Select role: 'Doctor'")
    add_bullet(doc, "Click 'Register' - you'll see 'Awaiting admin approval'")
    doc.add_paragraph()

    add_paragraph(doc, "Step 2: Wait for Admin Approval (24-48 hours)", bold=True)
    doc.add_paragraph()

    add_paragraph(doc, "Step 3: Login After Approval", bold=True)
    add_bullet(doc, "Go to http://localhost:3000/login")
    add_bullet(doc, "Enter your email and password")
    add_bullet(doc, "You'll be redirected to /dashboard/doctor")
    doc.add_paragraph()

    add_heading(doc, "5.3 Doctor Dashboard Overview", level=2)
    add_bullet(doc, "🏠 Dashboard - Overview and statistics")
    add_bullet(doc, "📅 Appointments - Manage patient appointments")
    add_bullet(doc, "📄 Medical Records - View records you've created")
    add_bullet(doc, "💬 Patient Feedback - View patient reviews and ratings")
    doc.add_paragraph()

    add_heading(doc, "5.4 Core Doctor Functions", level=2)

    add_heading(doc, "Function 1: Manage Appointments", level=3)
    add_paragraph(doc, "Accessing Appointments:")
    add_number(doc, 1, "Click 'Appointments' in sidebar")
    add_number(doc, 2, "View list of all patient appointments")
    doc.add_paragraph()

    add_paragraph(doc, "Appointment Actions:")
    add_bullet(doc, "Confirm Appointment: Review pending appointment and click 'Confirm'")
    add_bullet(doc, "Complete Appointment: After seeing patient, click 'Complete'")
    add_bullet(doc, "Add Medical Record: Click 'Add Record' to create patient record")
    doc.add_paragraph()

    add_heading(doc, "Function 2: Create Medical Records", level=3)
    add_paragraph(doc, "Medical Record Form includes:")
    add_bullet(doc, "Symptoms: Patient-reported symptoms and observations")
    add_bullet(doc, "Diagnosis: Medical diagnosis based on examination")
    add_bullet(doc, "Prescription: Medications with dosage instructions")
    add_bullet(doc, "Treatment Notes: Additional instructions and follow-up recommendations")
    doc.add_paragraph()

    add_paragraph(doc, "Example Medical Record:")
    add_paragraph(doc, "Symptoms: High fever (102°F), persistent cough, chest pain")
    add_paragraph(doc, "Diagnosis: Acute Bronchitis")
    add_paragraph(doc, "Prescription: Amoxicillin 500mg - 3 times daily for 7 days")
    add_paragraph(doc, "Treatment: Complete bed rest, plenty of fluids, follow-up in 7 days")
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # USER MANUAL - PATIENT
    add_heading(doc, "6. USER MANUAL - PATIENT", level=1)
    add_paragraph(doc, "Prepared by: Hillson (Frontend Developer & UI/UX Designer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "6.1 Patient Role Overview", level=2)
    add_paragraph(doc, "Patients use the Bayside HMS to:")
    add_bullet(doc, "Book appointments with doctors")
    add_bullet(doc, "View their medical records and diagnoses")
    add_bullet(doc, "Access prescriptions and treatment plans")
    add_bullet(doc, "Provide feedback to doctors")
    add_bullet(doc, "View and manage medical bills/invoices")
    doc.add_paragraph()

    add_heading(doc, "6.2 Patient Registration", level=2)
    add_paragraph(doc, "Step 1: Navigate to http://localhost:3000/register")
    doc.add_paragraph()
    
    add_paragraph(doc, "Step 2: Fill Personal Information", bold=True)
    add_bullet(doc, "First Name and Last Name")
    add_bullet(doc, "Email (used for login) and Password")
    add_bullet(doc, "Phone Number and Address")
    add_bullet(doc, "Gender and Date of Birth")
    add_bullet(doc, "Blood Group (A+, A-, B+, B-, AB+, AB-, O+, O-)")
    add_bullet(doc, "Emergency Contact Name and Phone")
    doc.add_paragraph()

    add_paragraph(doc, "Step 3: Select Role 'Patient' and Click 'Register'", bold=True)
    add_paragraph(doc, "Patient accounts are activated immediately (no admin approval needed)")
    doc.add_paragraph()

    add_heading(doc, "6.3 Core Patient Functions", level=2)

    add_heading(doc, "Function 1: Book an Appointment", level=3)
    add_paragraph(doc, "Complete Booking Process:")
    add_number(doc, 1, "Click 'Book Appointment' in sidebar")
    add_number(doc, 2, "Select Doctor from dropdown (shows specialization)")
    add_number(doc, 3, "Select Date using calendar widget (only future dates)")
    add_number(doc, 4, "Select Time from available slots (9 AM - 4 PM)")
    add_number(doc, 5, "Enter Reason for visit (be specific about symptoms)")
    add_number(doc, 6, "Click 'Book Appointment'")
    add_number(doc, 7, "Success message appears, appointment shows in 'Your Appointments'")
    doc.add_paragraph()

    add_heading(doc, "Function 2: View Medical Records", level=3)
    add_bullet(doc, "Click 'Medical Records' in sidebar")
    add_bullet(doc, "See complete medical history with all visits")
    add_bullet(doc, "Each record shows: Doctor, Date, Symptoms, Diagnosis, Prescription, Notes")
    add_bullet(doc, "Print or download records as PDF")
    add_bullet(doc, "Search and filter by date or doctor")
    doc.add_paragraph()

    add_heading(doc, "Function 3: Give Feedback to Doctor", level=3)
    add_bullet(doc, "Click 'Give Feedback' in sidebar")
    add_bullet(doc, "Select the doctor you want to review")
    add_bullet(doc, "Rate with 1-5 stars")
    add_bullet(doc, "Write detailed, honest feedback")
    add_bullet(doc, "Click 'Submit Feedback'")
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # ==========================================
    # NETWORKING CONFIGURATION
    # By: Niroj Shrestha
    # ==========================================
    add_heading(doc, "7. NETWORKING CONFIGURATION", level=1)
    add_paragraph(doc, "Prepared by: Niroj Shrestha (Networking and Testing Specialist)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "7.1 Network Architecture Overview", level=2)
    add_paragraph(doc, "The Bayside HMS uses a standard 3-tier web architecture:")
    add_bullet(doc, "Layer 1: User Devices (Clients) - Browsers, Mobile Apps")
    add_bullet(doc, "Layer 2: Web Server - Nginx/Apache reverse proxy")
    add_bullet(doc, "Layer 3: Application Server - Next.js (Port 3000)")
    add_bullet(doc, "Layer 4: Database Server - MySQL (Port 3306)")
    doc.add_paragraph()

    add_heading(doc, "7.2 Port Configuration", level=2)
    add_paragraph(doc, "Required Open Ports:", bold=True)
    doc.add_paragraph()
    add_paragraph(doc, "Production Server:")
    add_bullet(doc, "Port 80 - HTTP traffic (auto-redirect to HTTPS)")
    add_bullet(doc, "Port 443 - HTTPS encrypted traffic (primary)")
    add_bullet(doc, "Port 3000 - Next.js application (internal only)")
    add_bullet(doc, "Port 3306 - MySQL database (internal only)")
    add_bullet(doc, "Port 22 - SSH access (admin only)")
    doc.add_paragraph()

    add_heading(doc, "7.3 Firewall Configuration (Ubuntu UFW)", level=2)
    add_code_block(doc, """# Enable firewall
sudo ufw enable

# Allow SSH
sudo ufw allow 22/tcp

# Allow HTTP and HTTPS
sudo ufw allow 80/tcp
sudo ufw allow 443/tcp

# Deny external MySQL access
sudo ufw deny 3306/tcp

# Check status
sudo ufw status verbose""")
    doc.add_paragraph()

    add_heading(doc, "7.4 SSL/TLS Certificate Setup", level=2)
    add_paragraph(doc, "Using Let's Encrypt for Free SSL Certificates:", bold=True)
    add_code_block(doc, """# Install Certbot
sudo apt install certbot python3-certbot-nginx -y

# Obtain certificate
sudo certbot --nginx -d yourdomain.com

# Test auto-renewal
sudo certbot renew --dry-run""")
    doc.add_paragraph()

    add_heading(doc, "7.5 Database Network Security", level=2)
    add_paragraph(doc, "MySQL Configuration for Security:")
    add_code_block(doc, """# Edit: /etc/mysql/mysql.conf.d/mysqld.cnf

[mysqld]
# Bind to localhost only
bind-address = 127.0.0.1

# Connection limits
max_connections = 100
connect_timeout = 10""")
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # ==========================================
    # SECURITY TESTING
    # By: Niroj Shrestha
    # ==========================================
    add_heading(doc, "8. SECURITY TESTING AND IMPLEMENTATION", level=1)
    add_paragraph(doc, "Prepared by: Niroj Shrestha (Networking and Testing Specialist)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "8.1 Security Architecture Overview", level=2)
    add_paragraph(doc, "The Bayside HMS implements multiple layers of security:")
    add_bullet(doc, "Layer 1: Network Security - Firewall rules, SSL/TLS, DDoS protection")
    add_bullet(doc, "Layer 2: Application Security - HTTPS, security headers, input validation")
    add_bullet(doc, "Layer 3: Authentication - JWT tokens, RBAC, bcrypt hashing")
    add_bullet(doc, "Layer 4: Database Security - Parameterized queries, SQL injection prevention")
    doc.add_paragraph()

    add_heading(doc, "8.2 Authentication Security", level=2)

    add_heading(doc, "Password Security Implementation", level=3)
    add_paragraph(doc, "The system uses bcrypt with 10 salt rounds for password hashing.")
    doc.add_paragraph()
    add_paragraph(doc, "Password Requirements:")
    add_bullet(doc, "Minimum 8 characters")
    add_bullet(doc, "At least one uppercase letter")
    add_bullet(doc, "At least one lowercase letter")
    add_bullet(doc, "At least one digit")
    add_bullet(doc, "At least one special character (@$!%*?&)")
    doc.add_paragraph()

    add_heading(doc, "JWT Token Security", level=3)
    add_paragraph(doc, "Token Configuration:")
    add_bullet(doc, "Algorithm: HS256")
    add_bullet(doc, "Expiration: 24 hours")
    add_bullet(doc, "Storage: HTTP-only cookies (prevents XSS)")
    add_bullet(doc, "Secure flag enabled in production (HTTPS only)")
    add_bullet(doc, "SameSite attribute prevents CSRF")
    doc.add_paragraph()

    add_heading(doc, "8.3 SQL Injection Prevention", level=2)
    add_paragraph(doc, "All database queries use parameterized statements to prevent SQL injection.")
    doc.add_paragraph()
    add_paragraph(doc, "Example Safe Query:")
    add_code_block(doc, """const query = 'SELECT * FROM users WHERE email = ?';
const [rows] = await pool.execute(query, [email]);

// Even malicious input is treated as literal string""")
    doc.add_paragraph()

    add_heading(doc, "8.4 Cross-Site Scripting (XSS) Prevention", level=2)
    add_paragraph(doc, "Protections Implemented:")
    add_bullet(doc, "React automatically escapes all output")
    add_bullet(doc, "Content Security Policy (CSP) headers configured")
    add_bullet(doc, "No unsafe HTML rendering")
    add_bullet(doc, "All user input validated and sanitized")
    doc.add_paragraph()

    add_heading(doc, "8.5 Security Testing Procedures", level=2)

    add_heading(doc, "Test 1: SQL Injection Testing", level=3)
    add_paragraph(doc, "Test Cases:")
    add_bullet(doc, "Input: admin'-- | Expected: Login fails")
    add_bullet(doc, "Input: ' OR '1'='1 | Expected: No authentication bypass")
    add_bullet(doc, "Input: '; DROP TABLE users;-- | Expected: No tables dropped")
    doc.add_paragraph()

    add_heading(doc, "Test 2: XSS Testing", level=3)
    add_paragraph(doc, "Test Inputs:")
    add_bullet(doc, "<script>alert('XSS')</script> | Expected: Displayed as text")
    add_bullet(doc, "<img src=x onerror=\"alert\"> | Expected: No JavaScript execution")
    doc.add_paragraph()

    add_heading(doc, "Test 3: Authentication Testing", level=3)
    add_paragraph(doc, "Test Cases:")
    add_bullet(doc, "Access protected route without token | Expected: Redirect to login")
    add_bullet(doc, "Access admin route with patient token | Expected: Unauthorized")
    add_bullet(doc, "Use expired token | Expected: Redirect to login")
    doc.add_paragraph()

    add_heading(doc, "8.6 Security Audit Checklist", level=2)
    add_paragraph(doc, "Pre-Deployment Checklist:")
    add_bullet(doc, "☑ Passwords hashed with bcrypt")
    add_bullet(doc, "☑ JWT tokens properly verified")
    add_bullet(doc, "☑ RBAC on all protected routes")
    add_bullet(doc, "☑ All queries parameterized")
    add_bullet(doc, "☑ XSS prevention with output encoding")
    add_bullet(doc, "☑ CSRF prevention with SameSite cookies")
    add_bullet(doc, "☑ HTTPS enforced with valid SSL")
    add_bullet(doc, "☑ Security headers configured")
    add_bullet(doc, "☑ Database not externally accessible")
    doc.add_paragraph()

    # Add student contributions summary
    doc.add_page_break()
    add_heading(doc, "9. STUDENT CONTRIBUTIONS SUMMARY", level=1)
    doc.add_paragraph()

    add_heading(doc, "Krischal Jung Thakuri - Backend/Database Developer", level=2)
    add_paragraph(doc, "Responsibilities:")
    add_bullet(doc, "Database schema design and implementation")
    add_bullet(doc, "MySQL configuration and optimization")
    add_bullet(doc, "Backend API development")
    add_bullet(doc, "Database security implementation")
    doc.add_paragraph()
    add_paragraph(doc, "Sections Authored:")
    add_bullet(doc, "Section 3: Hardware and Software Setup")
    add_bullet(doc, "Database architecture and configuration")
    doc.add_paragraph()

    add_heading(doc, "Hillson - Frontend Developer & UI/UX Designer", level=2)
    add_paragraph(doc, "Responsibilities:")
    add_bullet(doc, "User interface design and implementation")
    add_bullet(doc, "React component development")
    add_bullet(doc, "Admin, Doctor, and Patient dashboard development")
    doc.add_paragraph()
    add_paragraph(doc, "Sections Authored:")
    add_bullet(doc, "Section 4: User Manual - Administrator")
    add_bullet(doc, "Section 5: User Manual - Doctor")
    add_bullet(doc, "Section 6: User Manual - Patient")
    doc.add_paragraph()

    add_heading(doc, "Niroj Shrestha - Networking and Testing Specialist", level=2)
    add_paragraph(doc, "Responsibilities:")
    add_bullet(doc, "Network architecture design")
    add_bullet(doc, "Security testing and implementation")
    add_bullet(doc, "Firewall and SSL/TLS configuration")
    add_bullet(doc, "Security vulnerability assessment")
    doc.add_paragraph()
    add_paragraph(doc, "Sections Authored:")
    add_bullet(doc, "Section 7: Networking Configuration")
    add_bullet(doc, "Section 8: Security Testing and Implementation")
    doc.add_paragraph()

    # Save the modified document
    print("Saving updated Final.docx...")
    doc.save('Final.docx')
    print("✅ Successfully updated Final.docx with all new sections!")
    print()
    print("Added sections:")
    print("  ✓ Section 3: Hardware and Software Setup (Krischal Jung Thakuri)")
    print("  ✓ Section 4: User Manual - Administrator (Hillson)")
    print("  ✓ Section 5: User Manual - Doctor (Hillson)")
    print("  ✓ Section 6: User Manual - Patient (Hillson)")
    print("  ✓ Section 7: Networking Configuration (Niroj Shrestha)")
    print("  ✓ Section 8: Security Testing (Niroj Shrestha)")
    print("  ✓ Section 9: Student Contributions Summary")

if __name__ == "__main__":
    main()
