#!/usr/bin/env python3
"""
Script to add comprehensive user manual sections to Final.docx
Organized by student contributions
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to the document"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    return p

def add_code_block(doc, code):
    """Add a code block to the document"""
    p = doc.add_paragraph(code, style='Intense Quote')
    return p

def main():
    # Open existing document
    print("Opening Final.docx...")
    doc = Document('Final.docx')

    print("Adding new sections...")

    # Add page break before new sections
    doc.add_page_break()

    # ==========================================
    # HARDWARE AND SOFTWARE SETUP
    # By: Krischal Jung Thakuri
    # ==========================================
    add_heading(doc, "3. HARDWARE AND SOFTWARE SETUP", level=1)
    add_paragraph(doc, "Prepared by: Krischal Jung Thakuri (Backend/Database Developer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "3.1 Hardware Requirements", level=2)

    add_heading(doc, "3.1.1 Development Environment", level=3)
    add_paragraph(doc, "Minimum Requirements:")
    doc.add_paragraph("Processor: Intel Core i3 / AMD Ryzen 3 or equivalent", style='List Bullet')
    doc.add_paragraph("RAM: 4GB minimum (8GB recommended)", style='List Bullet')
    doc.add_paragraph("Storage: 10GB free disk space", style='List Bullet')
    doc.add_paragraph("Network: Stable internet connection", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "3.1.2 Production Server Requirements", level=3)
    doc.add_paragraph("Processor: Intel Core i5 / AMD Ryzen 5 or better (quad-core)", style='List Bullet')
    doc.add_paragraph("RAM: 8GB minimum (16GB recommended for high traffic)", style='List Bullet')
    doc.add_paragraph("Storage: 50GB SSD (for database and application files)", style='List Bullet')
    doc.add_paragraph("Network: 100 Mbps internet connection with static IP", style='List Bullet')
    doc.add_paragraph("Backup Storage: External or cloud backup solution", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "3.2 Software Prerequisites", level=2)

    add_heading(doc, "3.2.1 Operating System", level=3)
    add_paragraph(doc, "The system supports multiple operating systems:")
    doc.add_paragraph("Windows 10/11", style='List Bullet')
    doc.add_paragraph("macOS 10.15 or higher", style='List Bullet')
    doc.add_paragraph("Linux (Ubuntu 20.04+ recommended)", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "3.2.2 Node.js and npm", level=3)
    add_paragraph(doc, "Node.js version 18.0.0 or higher is required.")
    add_code_block(doc, "# Verify installation:\nnode --version  # Should show v18.0.0 or higher\nnpm --version   # Should show 9.0.0 or higher")
    doc.add_paragraph()

    add_heading(doc, "3.2.3 MySQL Database Server", level=3)
    add_paragraph(doc, "MySQL 8.0 or higher is required for the database backend.")
    doc.add_paragraph()
    add_paragraph(doc, "Windows Installation:", bold=True)
    doc.add_paragraph("Download MySQL Installer from mysql.com", style='List Bullet')
    doc.add_paragraph("Run the installer and select 'MySQL Server' and 'MySQL Workbench'", style='List Bullet')
    doc.add_paragraph("Set root password during installation", style='List Bullet')
    doc.add_paragraph()

    add_paragraph(doc, "macOS Installation:", bold=True)
    add_code_block(doc, "brew install mysql\nbrew services start mysql\nmysql_secure_installation")
    doc.add_paragraph()

    add_paragraph(doc, "Linux (Ubuntu) Installation:", bold=True)
    add_code_block(doc, "sudo apt update\nsudo apt install mysql-server\nsudo systemctl start mysql\nsudo mysql_secure_installation")
    doc.add_paragraph()

    add_heading(doc, "3.3 Installation Steps", level=2)

    add_heading(doc, "Step 1: Clone the Repository", level=3)
    add_code_block(doc, "git clone https://github.com/Nischal1111/bayside-hms.git\ncd bayside-hms")

    add_heading(doc, "Step 2: Install Dependencies", level=3)
    add_code_block(doc, "npm install")
    add_paragraph(doc, "This installs all required packages including Next.js, React, database drivers, authentication libraries, and UI components.")
    doc.add_paragraph()

    add_heading(doc, "Step 3: Database Setup", level=3)
    add_paragraph(doc, "Create MySQL Database:", bold=True)
    add_code_block(doc, """mysql -u root -p

CREATE DATABASE bayside_hms CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;
CREATE USER 'bayside_admin'@'localhost' IDENTIFIED BY 'YourSecurePassword123!';
GRANT ALL PRIVILEGES ON bayside_hms.* TO 'bayside_admin'@'localhost';
FLUSH PRIVILEGES;
EXIT;""")
    doc.add_paragraph()

    add_paragraph(doc, "Import Database Schema:", bold=True)
    add_code_block(doc, "mysql -u bayside_admin -p bayside_hms < database/schema-mysql.sql")
    doc.add_paragraph()

    add_heading(doc, "Step 4: Environment Configuration", level=3)
    add_paragraph(doc, "Create .env.local file in the project root with the following content:")
    add_code_block(doc, """DATABASE_URL=mysql://bayside_admin:YourSecurePassword123!@localhost:3306/bayside_hms
JWT_SECRET=your_very_secure_random_32_character_secret_key_here
NODE_ENV=development
PORT=3000""")
    doc.add_paragraph()

    add_heading(doc, "Step 5: Create Admin Account", level=3)
    add_code_block(doc, "npm run create-admin")
    add_paragraph(doc, "Default Admin Credentials:")
    doc.add_paragraph("Email: admin@bayside-hms.com", style='List Bullet')
    doc.add_paragraph("Password: Admin@123456", style='List Bullet')
    add_paragraph(doc, "⚠️ IMPORTANT: Change this password immediately after first login!", bold=True)
    doc.add_paragraph()

    add_heading(doc, "Step 6: Start the Development Server", level=3)
    add_code_block(doc, "npm run dev")
    add_paragraph(doc, "The application will be available at: http://localhost:3000")
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # ==========================================
    # USER MANUALS
    # By: Hillson
    # ==========================================
    add_heading(doc, "4. USER MANUAL - ADMINISTRATOR", level=1)
    add_paragraph(doc, "Prepared by: Hillson (Frontend Developer & UI/UX Designer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "4.1 Administrator Role Overview", level=2)
    add_paragraph(doc, "Administrators are the system managers with the highest level of access. Their primary responsibilities include:")
    doc.add_paragraph("Approving new doctor registrations", style='List Bullet')
    doc.add_paragraph("Managing all users (patients and doctors)", style='List Bullet')
    doc.add_paragraph("Monitoring system statistics and operations", style='List Bullet')
    doc.add_paragraph("Overseeing appointment workflows", style='List Bullet')
    doc.add_paragraph("Ensuring system integrity and security", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "4.2 Accessing the Admin Portal", level=2)
    add_paragraph(doc, "Login Process:", bold=True)
    add_paragraph(doc, "Step 1: Navigate to Admin Portal")
    doc.add_paragraph("Open browser and go to: http://localhost:3000/admin", style='List Bullet')
    doc.add_paragraph("Or click 'Admin Portal' button on the homepage", style='List Bullet')
    doc.add_paragraph()

    add_paragraph(doc, "Step 2: Enter Credentials")
    doc.add_paragraph("Email: admin@bayside-hms.com", style='List Bullet')
    doc.add_paragraph("Password: Admin@123456 (change after first login)", style='List Bullet')
    doc.add_paragraph()

    add_paragraph(doc, "Step 3: Click 'Sign In'")
    doc.add_paragraph("You'll be redirected to the Admin Dashboard", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "4.3 Admin Dashboard Overview", level=2)
    add_paragraph(doc, "The admin dashboard features a sidebar with the following sections:")
    doc.add_paragraph("🏠 Dashboard - Main overview with statistics", style='List Bullet')
    doc.add_paragraph("👤 Pending Approvals - New doctor registrations awaiting approval", style='List Bullet')
    doc.add_paragraph("👨‍⚕️ Manage Doctors - All approved doctors", style='List Bullet')
    doc.add_paragraph("👥 Manage Patients - All registered patients", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "4.4 Core Administrator Functions", level=2)

    add_heading(doc, "Function 1: Approve Doctor Registrations", level=3)
    add_paragraph(doc, "Purpose: Verify and activate new doctor accounts")
    doc.add_paragraph()
    add_paragraph(doc, "Process:")
    doc.add_paragraph("1. Click 'Pending Approvals' in sidebar", style='List Number')
    doc.add_paragraph("2. Review doctor information (name, specialization, license number)", style='List Number')
    doc.add_paragraph("3. Verify credentials", style='List Number')
    doc.add_paragraph("4. Click green 'Approve' button", style='List Number')
    doc.add_paragraph("5. Confirm approval in dialog", style='List Number')
    doc.add_paragraph("6. Success message appears and doctor can now login", style='List Number')
    doc.add_paragraph()

    add_heading(doc, "Function 2: Manage Doctors", level=3)
    doc.add_paragraph("Click 'Manage Doctors' in sidebar", style='List Bullet')
    doc.add_paragraph("View list of all approved doctors with their specialization, phone, email", style='List Bullet')
    doc.add_paragraph("Click on doctor row to see complete profile and appointment history", style='List Bullet')
    doc.add_paragraph("Search by name, specialization, or license number", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Function 3: Manage Patients", level=3)
    doc.add_paragraph("Click 'Manage Patients' in sidebar", style='List Bullet')
    doc.add_paragraph("View list of all registered patients", style='List Bullet')
    doc.add_paragraph("See patient details including contact information, blood group, emergency contacts", style='List Bullet')
    doc.add_paragraph("View patient medical history and appointments", style='List Bullet')
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # USER MANUAL - DOCTOR
    add_heading(doc, "5. USER MANUAL - DOCTOR", level=1)
    add_paragraph(doc, "Prepared by: Hillson (Frontend Developer & UI/UX Designer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "5.1 Doctor Role Overview", level=2)
    add_paragraph(doc, "Doctors use the system to:")
    doc.add_paragraph("Manage patient appointments", style='List Bullet')
    doc.add_paragraph("Create and maintain electronic medical records (EMR)", style='List Bullet')
    doc.add_paragraph("View patient medical history", style='List Bullet')
    doc.add_paragraph("Prescribe medications and treatments", style='List Bullet')
    doc.add_paragraph("Receive patient feedback and ratings", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "5.2 Doctor Registration and Activation", level=2)
    add_paragraph(doc, "Step 1: Register for Account", bold=True)
    doc.add_paragraph("Navigate to http://localhost:3000/register", style='List Bullet')
    doc.add_paragraph("Fill in personal information (name, email, password)", style='List Bullet')
    doc.add_paragraph("Fill in professional information (phone, license number, specialization)", style='List Bullet')
    doc.add_paragraph("Select role: 'Doctor'", style='List Bullet')
    doc.add_paragraph("Click 'Register' - you'll see 'Awaiting admin approval'", style='List Bullet')
    doc.add_paragraph()

    add_paragraph(doc, "Step 2: Wait for Admin Approval", bold=True)
    add_paragraph(doc, "Your application is sent to hospital administrators")
    add_paragraph(doc, "Typical approval time: 24-48 hours")
    doc.add_paragraph()

    add_paragraph(doc, "Step 3: Login After Approval", bold=True)
    doc.add_paragraph("Go to http://localhost:3000/login", style='List Bullet')
    doc.add_paragraph("Enter your email and password", style='List Bullet')
    doc.add_paragraph("You'll be redirected to /dashboard/doctor", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "5.3 Doctor Dashboard Overview", level=2)
    doc.add_paragraph("🏠 Dashboard - Overview and statistics", style='List Bullet')
    doc.add_paragraph("📅 Appointments - Manage patient appointments", style='List Bullet')
    doc.add_paragraph("📄 Medical Records - View records you've created", style='List Bullet')
    doc.add_paragraph("💬 Patient Feedback - View patient reviews and ratings", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "5.4 Core Doctor Functions", level=2)

    add_heading(doc, "Function 1: Manage Appointments", level=3)
    add_paragraph(doc, "Accessing Appointments:")
    doc.add_paragraph("1. Click 'Appointments' in sidebar", style='List Number')
    doc.add_paragraph("2. View list of all patient appointments", style='List Number')
    doc.add_paragraph()

    add_paragraph(doc, "Appointment Actions:")
    doc.add_paragraph("Confirm Appointment: Review pending appointment and click 'Confirm' button", style='List Bullet')
    doc.add_paragraph("Complete Appointment: After seeing patient, click 'Complete' button", style='List Bullet')
    doc.add_paragraph("Add Medical Record: Click 'Add Record' to create patient medical record", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Function 2: Create Medical Records", level=3)
    add_paragraph(doc, "Medical Record Form includes:")
    doc.add_paragraph("Symptoms: Patient-reported symptoms and observations", style='List Bullet')
    doc.add_paragraph("Diagnosis: Medical diagnosis based on examination", style='List Bullet')
    doc.add_paragraph("Prescription: Medications prescribed with dosage instructions", style='List Bullet')
    doc.add_paragraph("Treatment Notes: Additional treatment instructions and follow-up recommendations", style='List Bullet')
    doc.add_paragraph()

    add_paragraph(doc, "Example Entry:")
    add_paragraph(doc, "Symptoms: High fever (102°F), persistent cough for 5 days, chest pain")
    add_paragraph(doc, "Diagnosis: Acute Bronchitis")
    add_paragraph(doc, "Prescription: Amoxicillin 500mg - 3 times daily for 7 days")
    add_paragraph(doc, "Treatment Notes: Complete bed rest, drink plenty of fluids, follow-up in 7 days")
    doc.add_paragraph()

    add_heading(doc, "Function 3: View Patient Feedback", level=3)
    doc.add_paragraph("Click 'Patient Feedback' in sidebar", style='List Bullet')
    doc.add_paragraph("View all feedback from patients you've treated", style='List Bullet')
    doc.add_paragraph("See ratings (1-5 stars) and written comments", style='List Bullet')
    doc.add_paragraph("Use feedback to improve service quality", style='List Bullet')
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # USER MANUAL - PATIENT
    add_heading(doc, "6. USER MANUAL - PATIENT", level=1)
    add_paragraph(doc, "Prepared by: Hillson (Frontend Developer & UI/UX Designer)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "6.1 Patient Role Overview", level=2)
    add_paragraph(doc, "Patients use the Bayside HMS to:")
    doc.add_paragraph("Book appointments with doctors", style='List Bullet')
    doc.add_paragraph("View their medical records and diagnoses", style='List Bullet')
    doc.add_paragraph("Access prescriptions and treatment plans", style='List Bullet')
    doc.add_paragraph("Provide feedback to doctors", style='List Bullet')
    doc.add_paragraph("View and manage medical bills/invoices", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "6.2 Patient Registration", level=2)
    add_paragraph(doc, "Step 1: Navigate to http://localhost:3000/register", bold=True)
    doc.add_paragraph()
    add_paragraph(doc, "Step 2: Fill Personal Information", bold=True)
    doc.add_paragraph("First Name and Last Name", style='List Bullet')
    doc.add_paragraph("Email (used for login) and Password", style='List Bullet')
    doc.add_paragraph("Phone Number and Address", style='List Bullet')
    doc.add_paragraph("Gender and Date of Birth", style='List Bullet')
    doc.add_paragraph("Blood Group (A+, A-, B+, B-, AB+, AB-, O+, O-)", style='List Bullet')
    doc.add_paragraph("Emergency Contact Name and Phone", style='List Bullet')
    doc.add_paragraph()

    add_paragraph(doc, "Step 3: Select Role 'Patient' and Click 'Register'", bold=True)
    add_paragraph(doc, "Patient accounts are activated immediately (no admin approval needed)")
    doc.add_paragraph()

    add_heading(doc, "6.3 Core Patient Functions", level=2)

    add_heading(doc, "Function 1: Book an Appointment", level=3)
    add_paragraph(doc, "Complete Booking Process:")
    doc.add_paragraph("1. Click 'Book Appointment' in sidebar", style='List Number')
    doc.add_paragraph("2. Select Doctor from dropdown (shows specialization)", style='List Number')
    doc.add_paragraph("3. Select Date using calendar widget (only future dates)", style='List Number')
    doc.add_paragraph("4. Select Time from available slots (9 AM - 4 PM)", style='List Number')
    doc.add_paragraph("5. Enter Reason for visit (be specific about symptoms)", style='List Number')
    doc.add_paragraph("6. Click 'Book Appointment'", style='List Number')
    doc.add_paragraph("7. Success message appears and appointment shows in 'Your Appointments'", style='List Number')
    doc.add_paragraph()

    add_heading(doc, "Function 2: View Medical Records", level=3)
    doc.add_paragraph("Click 'Medical Records' in sidebar", style='List Bullet')
    doc.add_paragraph("See complete medical history with all visits", style='List Bullet')
    doc.add_paragraph("Each record shows: Doctor, Date, Symptoms, Diagnosis, Prescription, Treatment Notes", style='List Bullet')
    doc.add_paragraph("Print or download records as PDF", style='List Bullet')
    doc.add_paragraph("Search and filter by date or doctor", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Function 3: Give Feedback to Doctor", level=3)
    doc.add_paragraph("Click 'Give Feedback' in sidebar", style='List Bullet')
    doc.add_paragraph("Select the doctor you want to review", style='List Bullet')
    doc.add_paragraph("Rate with 1-5 stars", style='List Bullet')
    doc.add_paragraph("Write detailed, honest feedback", style='List Bullet')
    doc.add_paragraph("Click 'Submit Feedback'", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Function 4: View Billing & Invoices", level=3)
    doc.add_paragraph("Click 'Billing & Invoices' in sidebar", style='List Bullet')
    doc.add_paragraph("View all medical bills and invoices", style='List Bullet')
    doc.add_paragraph("See payment status (Paid/Pending/Overdue)", style='List Bullet')
    doc.add_paragraph("Download or print invoices", style='List Bullet')
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # ==========================================
    # NETWORKING CONFIGURATION
    # By: Niroj Shrestha
    # ==========================================
    add_heading(doc, "7. NETWORKING CONFIGURATION", level=1)
    add_paragraph(doc, "Prepared by: Niroj Shrestha (Networking and Testing Specialist)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "7.1 Network Architecture Overview", level=2)
    add_paragraph(doc, "The Bayside HMS uses a standard 3-tier web architecture:")
    doc.add_paragraph("Layer 1: User Devices (Clients) - Browsers, Mobile Apps", style='List Bullet')
    doc.add_paragraph("Layer 2: Web Server - Nginx/Apache reverse proxy", style='List Bullet')
    doc.add_paragraph("Layer 3: Application Server - Next.js (Port 3000)", style='List Bullet')
    doc.add_paragraph("Layer 4: Database Server - MySQL (Port 3306)", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "7.2 Port Configuration", level=2)
    add_paragraph(doc, "Required Open Ports:", bold=True)
    doc.add_paragraph()
    add_paragraph(doc, "Production Server:")
    doc.add_paragraph("Port 80 - HTTP traffic (auto-redirect to HTTPS)", style='List Bullet')
    doc.add_paragraph("Port 443 - HTTPS encrypted traffic (primary)", style='List Bullet')
    doc.add_paragraph("Port 3000 - Next.js application (internal only)", style='List Bullet')
    doc.add_paragraph("Port 3306 - MySQL database (internal only, restrict external access)", style='List Bullet')
    doc.add_paragraph("Port 22 - SSH access (admin only, use key-based auth)", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "7.3 Firewall Configuration (Ubuntu UFW)", level=2)
    add_code_block(doc, """# Enable firewall
sudo ufw enable

# Allow SSH
sudo ufw allow 22/tcp

# Allow HTTP and HTTPS
sudo ufw allow 80/tcp
sudo ufw allow 443/tcp

# Deny external MySQL access
sudo ufw deny 3306/tcp

# Check status
sudo ufw status verbose""")
    doc.add_paragraph()

    add_heading(doc, "7.4 SSL/TLS Certificate Setup", level=2)
    add_paragraph(doc, "Using Let's Encrypt for Free SSL Certificates:", bold=True)
    add_code_block(doc, """# Install Certbot
sudo apt install certbot python3-certbot-nginx -y

# Obtain certificate
sudo certbot --nginx -d yourdomain.com -d www.yourdomain.com

# Auto-renewal is configured automatically
# Test renewal:
sudo certbot renew --dry-run""")
    doc.add_paragraph()

    add_heading(doc, "7.5 Nginx Configuration", level=2)
    add_paragraph(doc, "Production Nginx configuration with security headers:")
    add_code_block(doc, """# Redirect HTTP to HTTPS
server {
    listen 80;
    server_name bayside-hms.com;
    return 301 https://$server_name$request_uri;
}

# HTTPS Server
server {
    listen 443 ssl http2;
    server_name bayside-hms.com;

    # SSL Certificates
    ssl_certificate /etc/letsencrypt/live/bayside-hms.com/fullchain.pem;
    ssl_certificate_key /etc/letsencrypt/live/bayside-hms.com/privkey.pem;

    # Security Headers
    add_header X-Frame-Options "SAMEORIGIN" always;
    add_header X-Content-Type-Options "nosniff" always;
    add_header Strict-Transport-Security "max-age=31536000" always;

    # Proxy to Next.js
    location / {
        proxy_pass http://localhost:3000;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
    }
}""")
    doc.add_paragraph()

    add_heading(doc, "7.6 Database Network Security", level=2)
    add_paragraph(doc, "MySQL Configuration for Security:")
    add_code_block(doc, """# Edit: /etc/mysql/mysql.conf.d/mysqld.cnf

[mysqld]
# Bind to localhost only (prevent external access)
bind-address = 127.0.0.1

# Connection limits
max_connections = 100
connect_timeout = 10""")
    doc.add_paragraph()

    # Add page break
    doc.add_page_break()

    # ==========================================
    # SECURITY TESTING
    # By: Niroj Shrestha
    # ==========================================
    add_heading(doc, "8. SECURITY TESTING AND IMPLEMENTATION", level=1)
    add_paragraph(doc, "Prepared by: Niroj Shrestha (Networking and Testing Specialist)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "8.1 Security Architecture Overview", level=2)
    add_paragraph(doc, "The Bayside HMS implements multiple layers of security:")
    doc.add_paragraph("Layer 1: Network Security - Firewall rules, SSL/TLS encryption, DDoS protection", style='List Bullet')
    doc.add_paragraph("Layer 2: Application Security - HTTPS enforcement, security headers, input validation", style='List Bullet')
    doc.add_paragraph("Layer 3: Authentication & Authorization - JWT tokens, RBAC, bcrypt password hashing", style='List Bullet')
    doc.add_paragraph("Layer 4: Database Security - Parameterized queries, SQL injection prevention, encrypted connections", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "8.2 Authentication Security", level=2)

    add_heading(doc, "Password Security Implementation", level=3)
    add_paragraph(doc, "Bcrypt Hashing:")
    add_paragraph(doc, "The system uses bcrypt with 10 salt rounds for password hashing. Passwords are never stored in plain text.")
    doc.add_paragraph()

    add_paragraph(doc, "Password Requirements:")
    doc.add_paragraph("Minimum 8 characters", style='List Bullet')
    doc.add_paragraph("At least one uppercase letter", style='List Bullet')
    doc.add_paragraph("At least one lowercase letter", style='List Bullet')
    doc.add_paragraph("At least one digit", style='List Bullet')
    doc.add_paragraph("At least one special character (@$!%*?&)", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "JWT Token Security", level=3)
    add_paragraph(doc, "Token Configuration:")
    doc.add_paragraph("Algorithm: HS256", style='List Bullet')
    doc.add_paragraph("Expiration: 24 hours", style='List Bullet')
    doc.add_paragraph("Storage: HTTP-only cookies (prevents XSS attacks)", style='List Bullet')
    doc.add_paragraph("Secure flag enabled in production (HTTPS only)", style='List Bullet')
    doc.add_paragraph("SameSite attribute prevents CSRF attacks", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "8.3 SQL Injection Prevention", level=2)
    add_paragraph(doc, "All database queries use parameterized statements to prevent SQL injection attacks.")
    doc.add_paragraph()
    add_paragraph(doc, "Example Safe Query:")
    add_code_block(doc, """// ✅ SAFE - Parameterized query
const email = request.body.email;
const query = 'SELECT * FROM users WHERE email = ?';
const [rows] = await pool.execute(query, [email]);

// Even malicious input like "admin'--" is treated as literal string""")
    doc.add_paragraph()

    add_heading(doc, "8.4 Cross-Site Scripting (XSS) Prevention", level=2)
    add_paragraph(doc, "Protections Implemented:")
    doc.add_paragraph("React automatically escapes all output", style='List Bullet')
    doc.add_paragraph("Content Security Policy (CSP) headers configured", style='List Bullet')
    doc.add_paragraph("No use of dangerouslySetInnerHTML without sanitization", style='List Bullet')
    doc.add_paragraph("All user input is validated and sanitized", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "8.5 Cross-Site Request Forgery (CSRF) Prevention", level=2)
    add_paragraph(doc, "CSRF Protection Mechanisms:")
    doc.add_paragraph("SameSite cookie attribute set to 'strict'", style='List Bullet')
    doc.add_paragraph("State-changing requests use POST/PATCH/DELETE methods only", style='List Bullet')
    doc.add_paragraph("Token verification on all sensitive operations", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "8.6 Security Headers Implementation", level=2)
    add_paragraph(doc, "Comprehensive Security Headers (Nginx):")
    add_code_block(doc, """# Prevent clickjacking
add_header X-Frame-Options "SAMEORIGIN" always;

# Prevent MIME-type sniffing
add_header X-Content-Type-Options "nosniff" always;

# Enable XSS filter
add_header X-XSS-Protection "1; mode=block" always;

# HTTP Strict Transport Security
add_header Strict-Transport-Security "max-age=31536000; includeSubDomains" always;

# Content Security Policy
add_header Content-Security-Policy "default-src 'self'" always;""")
    doc.add_paragraph()

    add_heading(doc, "8.7 Security Testing Procedures", level=2)

    add_heading(doc, "Test 1: SQL Injection Testing", level=3)
    add_paragraph(doc, "Test Cases:")
    doc.add_paragraph("Input: admin'-- | Expected: Login fails, treated as literal string", style='List Bullet')
    doc.add_paragraph("Input: ' OR '1'='1 | Expected: Login fails, no authentication bypass", style='List Bullet')
    doc.add_paragraph("Input: '; DROP TABLE users;-- | Expected: No tables dropped, login fails", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Test 2: XSS Testing", level=3)
    add_paragraph(doc, "Test Inputs:")
    doc.add_paragraph("<script>alert('XSS')</script> | Expected: Displayed as text, not executed", style='List Bullet')
    doc.add_paragraph("<img src=x onerror=\"alert('XSS')\"> | Expected: No JavaScript execution", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Test 3: Authentication Bypass Testing", level=3)
    add_paragraph(doc, "Test Cases:")
    doc.add_paragraph("Access protected route without token | Expected: Redirect to login", style='List Bullet')
    doc.add_paragraph("Access admin route with patient token | Expected: Redirect to unauthorized", style='List Bullet')
    doc.add_paragraph("Use expired token | Expected: Redirect to login", style='List Bullet')
    doc.add_paragraph("Use tampered token | Expected: Redirect to login", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "8.8 Security Audit Checklist", level=2)
    add_paragraph(doc, "Pre-Deployment Security Checklist:")
    doc.add_paragraph("☑ Passwords hashed with bcrypt (10+ rounds)", style='List Bullet')
    doc.add_paragraph("☑ JWT tokens properly signed and verified", style='List Bullet')
    doc.add_paragraph("☑ RBAC implemented on all protected routes", style='List Bullet')
    doc.add_paragraph("☑ All SQL queries use parameterized statements", style='List Bullet')
    doc.add_paragraph("☑ XSS prevention with output encoding", style='List Bullet')
    doc.add_paragraph("☑ CSRF prevention with SameSite cookies", style='List Bullet')
    doc.add_paragraph("☑ HTTPS enforced with valid SSL certificate", style='List Bullet')
    doc.add_paragraph("☑ Security headers configured", style='List Bullet')
    doc.add_paragraph("☑ Database not accessible externally", style='List Bullet')
    doc.add_paragraph("☑ Regular backups configured", style='List Bullet')
    doc.add_paragraph()

    # Add final section
    doc.add_page_break()
    add_heading(doc, "9. STUDENT CONTRIBUTIONS SUMMARY", level=1)
    doc.add_paragraph()

    add_heading(doc, "Krischal Jung Thakuri - Backend/Database Developer", level=2)
    add_paragraph(doc, "Responsibilities:")
    doc.add_paragraph("Database schema design and implementation", style='List Bullet')
    doc.add_paragraph("MySQL configuration and optimization", style='List Bullet')
    doc.add_paragraph("Backend API development", style='List Bullet')
    doc.add_paragraph("Database security implementation", style='List Bullet')
    doc.add_paragraph("Backup and recovery procedures", style='List Bullet')
    doc.add_paragraph()
    add_paragraph(doc, "Sections Authored:")
    doc.add_paragraph("Section 3: Hardware and Software Setup", style='List Bullet')
    doc.add_paragraph("Database architecture and configuration", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Hillson - Frontend Developer & UI/UX Designer", level=2)
    add_paragraph(doc, "Responsibilities:")
    doc.add_paragraph("User interface design and implementation", style='List Bullet')
    doc.add_paragraph("React component development", style='List Bullet')
    doc.add_paragraph("Responsive design implementation", style='List Bullet')
    doc.add_paragraph("Admin, Doctor, and Patient dashboard development", style='List Bullet')
    doc.add_paragraph("Accessibility implementation", style='List Bullet')
    doc.add_paragraph()
    add_paragraph(doc, "Sections Authored:")
    doc.add_paragraph("Section 4: User Manual - Administrator", style='List Bullet')
    doc.add_paragraph("Section 5: User Manual - Doctor", style='List Bullet')
    doc.add_paragraph("Section 6: User Manual - Patient", style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, "Niroj Shrestha - Networking and Testing Specialist", level=2)
    add_paragraph(doc, "Responsibilities:")
    doc.add_paragraph("Network architecture design", style='List Bullet')
    doc.add_paragraph("Security testing and implementation", style='List Bullet')
    doc.add_paragraph("Firewall configuration", style='List Bullet')
    doc.add_paragraph("SSL/TLS setup", style='List Bullet')
    doc.add_paragraph("Security vulnerability assessment", style='List Bullet')
    doc.add_paragraph("Penetration testing", style='List Bullet')
    doc.add_paragraph()
    add_paragraph(doc, "Sections Authored:")
    doc.add_paragraph("Section 7: Networking Configuration", style='List Bullet')
    doc.add_paragraph("Section 8: Security Testing and Implementation", style='List Bullet')
    doc.add_paragraph()

    # Save the modified document
    print("Saving updated Final.docx...")
    doc.save('Final.docx')
    print("✅ Successfully updated Final.docx with all new sections!")
    print()
    print("Added sections:")
    print("  ✓ Hardware and Software Setup (Krischal Jung Thakuri)")
    print("  ✓ User Manual - Administrator (Hillson)")
    print("  ✓ User Manual - Doctor (Hillson)")
    print("  ✓ User Manual - Patient (Hillson)")
    print("  ✓ Networking Configuration (Niroj Shrestha)")
    print("  ✓ Security Testing and Implementation (Niroj Shrestha)")
    print("  ✓ Student Contributions Summary")

if __name__ == "__main__":
    main()
